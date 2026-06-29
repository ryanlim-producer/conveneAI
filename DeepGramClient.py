import os
import time
from deepgram import DeepgramClient, PrerecordedOptions
from fpdf import FPDF
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


class DeepgramPDFTranscriber:
    def __init__(self, api_key: str):
        if not api_key:
            raise ValueError("API key de Deepgram requerida.")
        self.client = DeepgramClient(api_key)

    def segundos_a_hhmmss(self, segundos: float) -> str:
        horas = int(segundos // 3600)
        minutos = int((segundos % 3600) // 60)
        segundos_restantes = int(segundos % 60)
        return f"[{horas:02d}:{minutos:02d}:{segundos_restantes:02d}]"

    def generar_pdf(self, nombre_salida: str, transcripciones: list[str]):
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", 'B', 14)
        pdf.cell(0, 10, "Transcripción de audio", ln=True, align="C")
        pdf.set_font("Arial", size=10)
        pdf.cell(0, 10, f"Archivo: {nombre_salida}", ln=True, align="C")
        pdf.ln(10)

        pdf.set_font("Arial", size=12)
        for linea in transcripciones:
            pdf.multi_cell(0, 10, linea)

        if not nombre_salida.lower().endswith(".pdf"):
            nombre_salida += ".pdf"

        pdf.output(nombre_salida)


    def transcribir_audio(self, ruta_audio, nombre_salida):
        inicio = time.time()
        try:
            if not ruta_audio or not os.path.isfile(ruta_audio):
                raise FileNotFoundError("❌ Archivo no válido o no encontrado.")

            if os.path.exists(nombre_salida):
                try:
                    with open(nombre_salida, "a"):
                        pass
                except PermissionError:
                    raise PermissionError(f"❌ El archivo '{nombre_salida}' está abierto en Word. Ciérralo e inténtalo de nuevo.")

            with open(ruta_audio, "rb") as f:
                audio_bytes = f.read()

            options = PrerecordedOptions(
                model="nova-2",
                language="es",
                smart_format=True,
                punctuate=True,
                paragraphs=True,
                diarize=True,
            )

            response = self.client.listen.prerecorded.v("1").transcribe_file(
                {"buffer": audio_bytes},
                options,
                timeout=1000
            )

            duracion_total = response.metadata.duration  # Duración total en segundos


            # Definir tamaño de bloque según duración
            if duracion_total <= 20 * 60:
                tam_bloque = 5 * 60
            elif duracion_total <= 60 * 60:
                tam_bloque = 15 * 60
            else:
                tam_bloque = 30 * 60

            channel = response.results.channels[0]
            transcripciones = []

            bloque_actual = -1
            for alt in channel.alternatives:
                if hasattr(alt, "paragraphs") and alt.paragraphs and hasattr(alt.paragraphs, "paragraphs"):
                    for paragraph in alt.paragraphs.paragraphs:
                        bloque_index = int(paragraph.start // tam_bloque)

                        # Si es un nuevo bloque, agregar encabezado
                        if bloque_index != bloque_actual:
                            bloque_actual = bloque_index
                            inicio_bloque = self.segundos_a_hhmmss(bloque_index * tam_bloque).strip("[]")
                            fin_bloque = self.segundos_a_hhmmss(
                                min((bloque_index + 1) * tam_bloque, duracion_total)
                            ).strip("[]")
                            transcripciones.append(f"\n==== Bloque {inicio_bloque} - {fin_bloque} ====\n")

                        tiempo = self.segundos_a_hhmmss(paragraph.start)
                        speaker = f"Locutor {paragraph.speaker}"
                        texto = " ".join([s.text for s in paragraph.sentences])
                        transcripciones.append(f"{speaker}: {texto.strip()}")

                elif alt.transcript and alt.transcript.strip():
                    transcripciones.append(alt.transcript.strip())

            if not transcripciones:
                raise ValueError("⚠ No se detectó voz en el archivo. Verifica que contenga audio hablado.")

            self.generar_word(nombre_salida, transcripciones, ruta_audio)

            fin = time.time()


        except Exception as e:
            print(f"❌ Exception: {e}")
            raise

    def obtener_fecha_creacion(self, ruta_archivo):
        try:
            fecha_creacion = os.path.getctime(ruta_archivo)
            fecha_str = f" Fecha de creación: {time.strftime('%Y-%m-%d %I:%M:%S %p', time.localtime(fecha_creacion))}"
            return fecha_str
        except Exception as e:
            return f" Fecha de creación (modificación): Error al obtener la fecha: {e}"

    def generar_word(self, nombre_salida: str, transcripciones: list[str], ruta_audio: str):
        doc = Document()

        titulo = doc.add_heading("Transcripción de audio", level=1)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        nombre_archivo = os.path.basename(ruta_audio)
        archivo_parrafo = doc.add_paragraph(f"Archivo: {nombre_archivo}")
        archivo_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        archivo_parrafo.runs[0].font.size = Pt(10)

        fecha_formateada = self.obtener_fecha_creacion(ruta_audio)
        fecha_parrafo = doc.add_paragraph(f"{fecha_formateada}")
        fecha_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fecha_parrafo.runs[0].font.size = Pt(10)

        doc.add_paragraph("")

        for linea in transcripciones:
            parrafo = doc.add_paragraph()
            run = parrafo.add_run(linea)
            run.font.size = Pt(12)

        if not nombre_salida.lower().endswith(".docx"):
            nombre_salida += ".docx"

        doc.save(nombre_salida)
